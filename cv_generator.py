from job_scrapper import (compare_social_to_database,
                          get_user_preferences,
                          get_user_projects_with_readmes)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (Paragraph, SimpleDocTemplate,
                                Spacer, ListFlowable, ListItem)
from reportlab.lib.pagesizes import letter
from reportlab.lib.enums import TA_CENTER, TA_LEFT

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from io import BytesIO
from langchain.schema.output_parser import StrOutputParser
from langchain_core.runnables import Runnable
from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq
import os
from dotenv import load_dotenv

load_dotenv()
CV_GENERATOR_GROQ_API_KEY = os.environ.get("CV_GENERATOR_GROQ_API_KEY")

cv_generator_llm = ChatGroq(
    model="llama-3.3-70b-versatile",
    temperature=0.3,
    max_tokens=1000,
    timeout=30,
    max_retries=2,
    api_key=CV_GENERATOR_GROQ_API_KEY
)

def get_user_full_details(user_id):
    messages = []
    state = {"llm_limit_exceeded": False}

    user_preferences = get_user_preferences(user_id)
    if not user_preferences:
        messages.append("User preferences not found.")
        return None, messages

    github_url = user_preferences.get("github")
    if not github_url:
        messages.append("GitHub URL not provided.")
        return user_preferences, messages

    github_username = github_url.split("/")[-1].strip()
    if not github_username:
        messages.append("Invalid GitHub URL.")
        return user_preferences, messages

    user_github_repo = get_user_projects_with_readmes(github_username)

    full_details = compare_social_to_database(user_preferences, user_github_repo, state)
    if not full_details:
        messages.append("Failed to merge GitHub data with user profile.")
        return user_preferences, messages

    return full_details, messages


def generate_cv_with_llm(data, follow_format=False):
    prompt_template = ChatPromptTemplate.from_messages([
        ("system",
        "You are a professional ATS-friendly CV generator.\n\n"
        "Instructions:\n"
        "1. If `follow_format` is True:\n"
        "   - Follow the original CV structure provided in `initial_cv_content`.\n"
        "2. If `follow_format` is False:\n"
        "   - Ignore the original format and create a new CV from scratch.\n"
        "3. Only include content relevant to the job and the user's capacity.\n"
        "4. Your response must ONLY contain the final CV content (no preamble, no extra explanation),\n"
        "   as it will be converted directly to PDF."
        ),
        ("user",
        "Here is the input data:\n\n"
        "{data}\n\n"
        "Follow original format: {follow_format}")
    ])

    
    chain: Runnable = prompt_template | cv_generator_llm | StrOutputParser()
    try:
        response = chain.invoke({
            "data": data,
            "follow_format": follow_format
        }).strip()

        return response
    
    except Exception as e:
        print(f"LLM failed to generate CV: {e}")
        return False

def convert_text_to_pdf(text):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=40, leftMargin=40,
                            topMargin=40, bottomMargin=40)

    styles = getSampleStyleSheet()
    
    # Custom styles
    styles.add(ParagraphStyle(name='Header', fontSize=14, leading=16, spaceAfter=10, spaceBefore=12, alignment=TA_LEFT, fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='ListItem', leftIndent=15, bulletIndent=5, spaceAfter=5, fontSize=10))

    flowables = []
    paragraphs = text.strip().split("\n")

    bullet_items = []

    for line in paragraphs:
        line = line.strip()

        if not line:
            continue  # skip empty lines

        # Detect headers like "# Summary"
        if line.startswith("#"):
            # Flush any previous bullet list before header
            if bullet_items:
                flowables.append(ListFlowable(bullet_items, bulletType='bullet', start='circle'))
                flowables.append(Spacer(1, 8))
                bullet_items = []

            header_text = line.lstrip("#").strip()
            flowables.append(Paragraph(header_text, styles['Header']))

        # Detect bullet list items like "- Python"
        elif line.startswith("- "):
            bullet_text = line[2:].strip()
            bullet_items.append(ListItem(Paragraph(bullet_text, styles["ListItem"])))

        else:
            # Flush any bullet items first
            if bullet_items:
                flowables.append(ListFlowable(bullet_items, bulletType='bullet', start='circle'))
                flowables.append(Spacer(1, 8))
                bullet_items = []

            flowables.append(Paragraph(line, styles["Normal"]))
            flowables.append(Spacer(1, 6))

    # Final bullet list (if any)
    if bullet_items:
        flowables.append(ListFlowable(bullet_items, bulletType='bullet', start='circle'))

    doc.build(flowables)
    return buffer.getvalue()


def convert_text_to_docx(text):
    doc = Document()
    buffer = BytesIO()

    # Define basic styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = text.strip().split('\n')
    in_list = False

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Section Headers
        if line.startswith('#'):
            if in_list:
                in_list = False  # End bullet list

            heading = line.lstrip('#').strip()
            para = doc.add_paragraph(heading)
            para.style = 'Heading 2'
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Bullet points
        elif line.startswith('- '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
            in_list = True

        # Regular paragraph
        else:
            if in_list:
                in_list = False  # End bullet list
            doc.add_paragraph(line)

    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()